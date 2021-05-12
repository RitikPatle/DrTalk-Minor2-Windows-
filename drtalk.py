from tkinter import *
import sqlite3
import pyttsx3
import speech_recognition as sr
import os,sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm , RGBColor , Inches
from docx.shared import Pt 
import datetime
import re
from tkinter import messagebox

window = Tk()
window.title('DrTalk')
window.wm_state('zoomed')                       #to open window in maximize mode
window.resizable(0,0)                           #to disable the Restore Down button
window.configure(background='#ffffff')          #to provide a background

iconvar = PhotoImage(file='icon.gif')           #to give icon
window.iconphoto(True,iconvar)                  #to give icon

logovar = PhotoImage(file='logo.gif')           #to use logo
logo = Label(window,image=logovar)              #to use logo
logo.place(x=-2,y=0)                            #to use logo

def SpeakText(sentence):
    engine = pyttsx3.init()
    engine.say(sentence)
    engine.runAndWait()

def ListenText():
    r = sr.Recognizer()
    with sr.Microphone() as source:
        print('I am listening')
        r.adjust_for_ambient_noise(source,duration=1)
        audio = r.listen(source)
        try:
            lisningvar = r.recognize_google(audio)
        except sr.UnknownValueError:
            lisningvar = ''
    return lisningvar

def cnvp(NaMe,GeNder,DoBi,PhNu,eID,REGnO,ExPerI,ClNc,DGre,UsRnme,PaSSwD,AoNE,AtWO,AtRE):
    pname = StringVar()
    pgendr = StringVar()
    page = StringVar()
    sympt = StringVar()
    diag = StringVar()
    pres = StringVar()
    adv = StringVar()
    
    def storeprescrip():
        prescripdoc = Document()

        RESC=prescripdoc.add_paragraph("")
        RESC.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cLNC=RESC.add_run(f"""{ClNc}\n""")
        cLNC.font.size=Pt(20)
        cLNC.font.color.rgb=RGBColor(0,111,80)

        pNAM=RESC.add_run("""Patient's Name: """)
        pNAM.font.size=Pt(16)
        #pNAM.font.color.rgb=RGBColor(0,0,0)

        pNAME=RESC.add_run(f"""{pname.get()}\n""")
        pNAME.font.size=Pt(16)
        pNAME.font.color.rgb=RGBColor(0,111,80)

        pGEN=RESC.add_run("""Gender: """)
        pGEN.font.size=Pt(16)
        #pGEN.font.color.rgb=RGBColor(0,0,0)

        pGEND=RESC.add_run(f"""{pgendr.get()}\t\t""")
        pGEND.font.size=Pt(16)
        pGEND.font.color.rgb=RGBColor(0,111,80)

        pAG=RESC.add_run("""Age: """)
        pAG.font.size=Pt(16)
        #pAG.font.color.rgb=RGBColor(0,0,0)

        pAGE=RESC.add_run(f"""{page.get()}""")
        pAGE.font.size=Pt(16)
        pAGE.font.color.rgb=RGBColor(0,111,80)

        pRESC=prescripdoc.add_paragraph("")
        pRESC.alignment = WD_ALIGN_PARAGRAPH.LEFT

        sYMPH=pRESC.add_run("""Symptoms:\n""")
        sYMPH.font.size=Pt(16)
        #sYMPH.font.color.rgb=RGBColor(0,0,0)

        sYMPT=pRESC.add_run(f"""{SymptomS.get("1.0",'end-1c')}\n""")
        sYMPT.font.size=Pt(14)
        sYMPT.font.color.rgb=RGBColor(0,111,80)

        dIAGH=pRESC.add_run("""Diagnosis:\n""")
        dIAGH.font.size=Pt(16)
        #dIAGH.font.color.rgb=RGBColor(0,0,0)

        dIAGT=pRESC.add_run(f"""{DiagnosiS.get("1.0",'end-1c')}\n""")
        dIAGT.font.size=Pt(14)
        dIAGT.font.color.rgb=RGBColor(0,111,80)

        pRESH=pRESC.add_run("""Prescription:\n""")
        pRESH.font.size=Pt(16)
        #pRESH.font.color.rgb=RGBColor(0,0,0)

        pREST=pRESC.add_run(f"""{PrescriptioN.get("1.0",'end-1c')}\n""")
        pREST.font.size=Pt(14)
        pREST.font.color.rgb=RGBColor(0,111,80)

        aDVCH=pRESC.add_run("""Advice:\n""")
        aDVCH.font.size=Pt(16)
        #aDVCH.font.color.rgb=RGBColor(0,0,0)

        aDVCT=pRESC.add_run(f"""{AdvicE.get("1.0",'end-1c')}\n""")
        aDVCT.font.size=Pt(14)
        aDVCT.font.color.rgb=RGBColor(0,111,80)

        pLABEL=pRESC.add_run("""Prescribed by,\n""")
        pLABEL.font.size=Pt(16)
        #pLABEL.font.color.rgb=RGBColor(0,0,0)

        dNAME=pRESC.add_run(f"""{NaMe}\n""")
        dNAME.font.size=Pt(16)
        dNAME.font.color.rgb=RGBColor(0,111,80)

        dRGN=pRESC.add_run("""Registration No.: """)
        dRGN.font.size=Pt(14)
        #dRGN.font.color.rgb=RGBColor(0,0,0)

        dRGNO=pRESC.add_run(f"""{REGnO}\n""")
        dRGNO.font.size=Pt(14)
        dRGNO.font.color.rgb=RGBColor(0,111,80)

        dDGRE=pRESC.add_run(f"""{DGre}\n""")
        dDGRE.font.size=Pt(14)
        dDGRE.font.color.rgb=RGBColor(0,111,80)

        dEX=pRESC.add_run("""Experience: """)
        dEX.font.size=Pt(14)
        #dEX.font.color.rgb=RGBColor(0,0,0)

        dEXP=pRESC.add_run(f"""{ExPerI}y\t""")
        dEXP.font.size=Pt(14)
        dEXP.font.color.rgb=RGBColor(0,111,80)

        dPH=pRESC.add_run("""Phone No.: """)
        dPH.font.size=Pt(14)
        #dPH.font.color.rgb=RGBColor(0,0,0)

        dPHN=pRESC.add_run(f"""{PhNu}\n""")
        dPHN.font.size=Pt(14)
        dPHN.font.color.rgb=RGBColor(0,111,80)

        dEI=pRESC.add_run("""Email Id: """)
        dEI.font.size=Pt(14)
        #dEI.font.color.rgb=RGBColor(0,0,0)

        dEID=pRESC.add_run(f"""{eID}\n""")
        dEID.font.size=Pt(14)
        dEID.font.color.rgb=RGBColor(00,111,80)

        dAT=pRESC.add_run("""On Date: """)
        dAT.font.size=Pt(14)
        #dAT.font.color.rgb=RGBColor(0,0,0)

        dATE=pRESC.add_run(f"""{datetime.datetime.now().strftime('%d/%b/%Y')}\t""")
        dATE.font.size=Pt(14)
        dATE.font.color.rgb=RGBColor(00,111,80)

        tIM=pRESC.add_run("""At Time: """)
        tIM.font.size=Pt(14)
        #tIM.font.color.rgb=RGBColor(0,0,0)

        tIME=pRESC.add_run(f"""{datetime.datetime.now().strftime('%I:%M:%S %p')}""")
        tIME.font.size=Pt(14)
        tIME.font.color.rgb=RGBColor(00,111,80)

        currentdtinfo = datetime.datetime.now()
        docdtinfo = currentdtinfo.strftime('%Y%b%d%H%M%S')
        prscdocnme = f'{docdtinfo}{pname.get()}'
        prescripdoc.save(f'Prescriptions\{prscdocnme}.docx')
        
        pdb = sqlite3.connect('drtalk')
        pcr = pdb.cursor()
        pcr.execute("""create table if not exists ps(Name VARCHAR,Gender VARCHAR,Age VARCHAR,Symptoms VARCHAR,Diagnosis VARCHAR,Prescription VARCHAR,Advice VARCHAR,DrregNO VARCHAR,Prscrpnme VARCHAR)""")
        pcr.execute("""insert into ps VALUES(?,?,?,?,?,?,?,?,?)""",(pname.get(), pgendr.get(), page.get(), SymptomS.get("1.0",'end-1c'), DiagnosiS.get("1.0",'end-1c'), PrescriptioN.get("1.0",'end-1c'), AdvicE.get("1.0",'end-1c'),REGnO,prscdocnme))
        pdb.commit()
        pdb.close()
        VpLisnwndw.destroy()
            
    def lISten():
        SpeakText('Name')
        sname = ListenText()
        pName.insert(0,sname)

        SpeakText('Gender')
        srgndr = ListenText()
        pGender.insert(0,srgndr)

        SpeakText('Age')
        srage = ListenText()
        pAge.insert(0,srage)

        SpeakText('Symptoms')
        srsympt = ListenText()
        SymptomS.insert(0.0,srsympt)

        SpeakText('Diagnosis')
        srdiag = ListenText()
        DiagnosiS.insert(0.0,srdiag)

        SpeakText('Prescription')
        srpresc = ListenText()
        PrescriptioN.insert(0.0,srpresc)

        SpeakText('Advice')
        sradv = ListenText()
        AdvicE.insert(0.0,sradv)

    VpLisnwndw = Toplevel(bg='#ffffff')
    VpLisnwndw.wm_state('zoomed')
    VpLisnwndw.resizable(0,0)

    pInfO = LabelFrame(VpLisnwndw,text='Patient Information:',font='-size 18 -weight bold',bd=6,bg='#ffffff',height=125,width=1535)
    pInfO.place(x=0,y=5)

    pNamE = Label(pInfO,background='white',text='Name:',font='-size 18')
    pNamE.place(x=5,y=5)

    pName = Entry(pInfO,background='#ffffff',textvariable=pname,font='-size 17',width=110)
    pName.place(x=85,y=7)

    pGendeR = Label(pInfO,background='white',text='Gender:',font='-size 18')
    pGendeR.place(x=5,y=45)

    pGender = Entry(pInfO,background='#ffffff',textvariable=pgendr,font='-size 17',width=10)
    pGender.place(x=105,y=47)

    pAgE = Label(pInfO,background='white',text='Age:',font='-size 18')
    pAgE.place(x=265,y=45)

    pAge = Entry(pInfO,background='#ffffff',textvariable=page,font='-size 17',width=3)
    pAge.place(x=325,y=47)

    SymPtomS = Label(VpLisnwndw,background='white',text='Symptoms:',font='-size 18')
    SymPtomS.place(x=5,y=135)

    SymptomS = Text(VpLisnwndw,font='-size 17',bg='#ffffff',height=5,width=117)
    SymptomS.place(x=5,y=170)

    DiaGnosiS = Label(VpLisnwndw,background='white',text='Diagnosis:',font='-size 18')
    DiaGnosiS.place(x=5,y=305)
    
    DiagnosiS = Text(VpLisnwndw,font='-size 17',bg='#ffffff',height=5,width=117)
    DiagnosiS.place(x=5,y=338)

    PreScriptioN = Label(VpLisnwndw,background='white',text='Prescription:',font='-size 18')
    PreScriptioN.place(x=5,y=470)

    PrescriptioN = Text(VpLisnwndw,font='-size 17',bg='#ffffff',height=5,width=117)
    PrescriptioN.place(x=5,y=500)

    AdVicE = Label(VpLisnwndw,background='white',text='Advice:',font='-size 18')
    AdVicE.place(x=5,y=634)

    AdvicE = Text(VpLisnwndw,font='-size 17',bg='#ffffff',height=5,width=117)
    AdvicE.place(x=5,y=662)

    SaVeButton = Button(VpLisnwndw,text='Save',cursor='hand2',background='#006F50',foreground='#ffffff',font='-size 15',height=0,width=8,command=storeprescrip)
    SaVeButton.place(x=5,y=800)

    checkgoogle = os.system('ping -n 1 www.google.com | FIND "Reply"')
    if checkgoogle==0:
        VpLisnwndw.after(2000,lISten)

def ChangE(cNme,cGdr,cDob,cPhn,cEid,cReg,cExp,cCln,cDgr,cUsr,cPas,cAon,cAtw,cAth):
    fname = StringVar()
    gendr = StringVar()
    dob = StringVar()
    phoneno = StringVar()
    eid = StringVar()
    regno = StringVar()
    exp = StringVar()
    cln = StringVar()
    dgr = StringVar()
    uname = StringVar()
    passwd = StringVar()
    aone = StringVar()
    atwo = StringVar()
    athr = StringVar()

    docdb = sqlite3.connect('drtalk')
    drcr = docdb.cursor()
    drcr.execute("""select * from drs where Uname=?""",[cUsr])
    res=drcr.fetchall()
    docdb.commit()
    docdb.close()
    for row in res:
        NAME=row[0]
        GEND=row[1]
        DOB=row[2]
        PHNO=row[3]
        EID=row[4]
        REGN=row[5]
        EXP=row[6]
        CLN=row[7]
        DGR=row[8]
        UNME=row[9]
        PASS=row[10]
        AONE=row[11]
        ATWO=row[12]
        ATHR=row[13]

    def changerec():
        nmatch=0
        gmatch=0
        dobmatch=0
        phnomatch=0
        eidmatch=0
        regnomatch=0
        expmatch=0
        dgrmatch=0
        passwdmatch=0
        aonematch=0
        atwomatch=0
        athrmatch=0
        NamePattern=r'[a-zA-Z]|\.'
        if re.match(NamePattern,fname.get()):
            nmatch=1
        else:
            nmatch=0
        GenderPattern=r'Male|Female|Others'
        if re.match(GenderPattern,gendr.get()):
            gmatch=1
        else:
            gmatch=0
        DobPattern=r'\d{2}/\d{2}/\d{4}'
        if re.match(DobPattern,dob.get()):
            dobmatch=1
        else:
            dobmatch=0
        PhnoPattern=r'\d{10}'
        if re.match(PhnoPattern,phoneno.get()):
            phnomatch=1
        else:
            phnomatch=0
        EidPattern=r'[a-zA-Z0-9.]*@[a-zA-Z]*\.com'
        if re.match(EidPattern,eid.get()):
            eidmatch=1
        else:
            eidmatch=0
        RegnoPattern=r'\d{6}'
        if re.match(RegnoPattern,regno.get()):
            regnomatch=1
        else:
            regnomatch=0
        ExpPattern=r'\d{2}'
        if re.match(ExpPattern,exp.get()):
            expmatch=1
        else:
            expmatch=0
        DgrPattern=r'[a-zA-Z0-9. -]'
        if re.match(DgrPattern,dgr.get()):
            dgrmatch=1
        else:
            dgrmatch=0
        PasswdPattern=r'(?=.*\d)(?=.*[a-z])(?=.*[A-Z])(?=.*\W)'
        if re.match(PasswdPattern,passwd.get()):
            passwdmatch=1
        else:
            passwdmatch=0
        AnsPattern=r'[A-Z]+'
        if re.match(AnsPattern,aone.get()):
            aonematch=1
        else:
            aonematch=0
        if re.match(AnsPattern,atwo.get()):
            atwomatch=1
        else:
            atwomatch=0
        if re.match(AnsPattern,athr.get()):
            athrmatch=1
        else:
            athrmatch=0
        if nmatch and gmatch and dobmatch and phnomatch and eidmatch and regnomatch and expmatch and dgrmatch and passwdmatch and aonematch and atwomatch and athrmatch==1:
            print('Matched')
            docdb = sqlite3.connect('drtalk')
            drcr = docdb.cursor()
            drcr.execute("""UPDATE drs SET Name=?,Gender=?,Dob=?,PhoneNo=?,Eid=?,RegNo=?,Exp=?,Cln=?,Dgr=?,Uname=?,Passwd=?,Aone=?,Atwo=?,Athr=? where Uname=?""",[fname.get(), gendr.get(), dob.get(), phoneno.get(), eid.get(), regno.get(), exp.get(), cln.get(), dgr.get(), uname.get(), passwd.get(), aone.get(), atwo.get(), athr.get(),cUsr])
            docdb.commit()
            docdb.close()
            window.destroy()
        else:
            print('Not Matched')
            messagebox.showinfo("Wrong Entries", "Please Enter Informations in correct format")

    def gOback():
        loggedin(cNme,cGdr,cDob,cPhn,cEid,cReg,cExp,cCln,cDgr,cUsr,cPas,cAon,cAtw,cAth)
    
    changeframe = Frame(window,bg='#ffffff',height=649,width=1536,background='#ffffff')
    changeframe.place(x=0,y=191)
    
    personalinfo = LabelFrame(changeframe,text='Personal Information:',font='-size 18 -weight bold',bd=6,bg='#ffffff',height=130,width=1535)
    personalinfo.place(x=0,y=5)

    nam = Label(personalinfo,text='Fullname:',bg='#ffffff',font='-size 18')
    nam.place(x=5,y=5)
    nme = Entry(personalinfo,bg='#ffffff',textvariable=fname,font='-size 17',width=107)
    nme.place(x=119,y=7)
    nme.insert(0,NAME)

    gndr = Label(personalinfo,text='Gender:',bg='#ffffff',font='-size 18')
    gendrbtn1 = Radiobutton(personalinfo,text='Male',tristatevalue='x',bg='#ffffff',variable=gendr,value='Male',font='-size 16')
    gendrbtn2 = Radiobutton(personalinfo,text='Female',tristatevalue='x',bg='#ffffff',variable=gendr,value='Female',font='-size 16')
    gendrbtn3 = Radiobutton(personalinfo,text='Others',tristatevalue='x',bg='#ffffff',variable=gendr,value='Others',font='-size 16')
    gndr.place(x=5,y=45)
    gendrbtn1.place(x=100,y=44)
    gendrbtn2.place(x=180,y=44)
    gendrbtn3.place(x=286,y=44)
    if GEND=='Male':
        gendrbtn1.select()
    elif GEND=='Female':
        gendrbtn2.select()
    elif GEND=='Others':
        gendrbtn3.select()
    else:
        print('Wrong value')

    DoB = Label(personalinfo,text='D.O.B (dd/mm/yyyy):',bg='#ffffff',font='-size 18')
    DoB.place(x=395,y=45)
    Dob = Entry(personalinfo,bg='#ffffff',textvariable=dob,font='-size 17',width=9)
    Dob.place(x=628,y=47)
    Dob.insert(0,DOB)

    Phn = Label(personalinfo,text='Contact Number:',bg='#ffffff',font='-size 18')
    Phn.place(x=760,y=45)
    Phno = Entry(personalinfo,bg='#ffffff',textvariable=phoneno,font='-size 17',width=10)
    Phno.place(x=950,y=47)
    Phno.insert(0,PHNO)

    Eml = Label(personalinfo,text='Email Id:',bg='#ffffff',font='-size 18')
    Eml.place(x=1110,y=45)
    Email = Entry(personalinfo,bg='#ffffff',textvariable=eid,font='-size 17',width=23)
    Email.place(x=1210,y=47)
    Email.insert(0,EID)

    professionalinfo = LabelFrame(changeframe,text='Professional Information:',font='-size 18 -weight bold',bd=6,bg='#ffffff',height=130,width=1535)
    professionalinfo.place(x=0,y=140)

    Regno = Label(professionalinfo,text='Registeration Number:',bg='#ffffff',font='-size 18')
    Regno.place(x=0,y=5)
    RegNo = Entry(professionalinfo,bg='#ffffff',textvariable=regno,font='-size 17',width=6)
    RegNo.place(x=250,y=7)
    RegNo.insert(0,REGN)

    ExP = Label(professionalinfo,text='Experience (In Years):',bg='#ffffff',font='-size 18')
    ExP.place(x=350,y=5)
    Exp = Entry(professionalinfo,bg='#ffffff',textvariable=exp,font='-size 17',width=2)
    Exp.place(x=595,y=7)
    Exp.insert(0,EXP)

    ClN = Label(professionalinfo,text='Clinic Name:',bg='#ffffff',font='-size 18')
    ClN.place(x=644,y=5)
    Cln = Entry(professionalinfo,bg='#ffffff',textvariable=cln,font='-size 17',width=55)
    Cln.place(x=793,y=7)
    Cln.insert(0,CLN)

    DgR = Label(professionalinfo,text='Education Qualification:',bg='#ffffff',font='-size 18')
    DgR.place(x=0,y=45)
    Dgr = Entry(professionalinfo,bg='#ffffff',textvariable=dgr,font='-size 17',width=96)
    Dgr.place(x=260,y=47)
    Dgr.insert(0,DGR)

    UsrNme = Label(changeframe,text='Username:',font='-size 18 -weight bold',bg='#ffffff')
    UsrNme.place(x=4,y=270)
    usernam = Entry(changeframe,background='#ffffff',textvariable=uname,font='-size 17',width=48)
    usernam.place(x=130,y=274)
    usernam.insert(0,UNME)

    pas = Label(changeframe,background='white',text='Password:',font='-size 18 -weight bold')
    pas.place(x=760,y=270)
    psswd = Entry(changeframe,background='#ffffff',textvariable=passwd,font='-size 17',width=48)
    psswd.place(x=890,y=274)
    psswd.insert(0,PASS)

    tac = LabelFrame(changeframe,text='Terms and Condition:',font='-size 18 -weight bold',bd=6,bg='#ffffff',height=285,width=1535)
    tac.place(x=0,y=307)

    fpaes1 = Label(tac,text='By chance if you forgot your password then by clicking on forgot password, You have answer the below given questions and you can change',bg='#ffffff',font='-size 18')
    fpaes1.place(x=0,y=5)

    fpaes2 = Label(tac,text='your password. Another thing is the application will accept only those Email Ids which are based on .com servers so you must enter only tho-',bg='#ffffff',font='-size 18')
    fpaes2.place(x=0,y=45)

    fpaes3 = Label(tac,text='-se emails having .com at end. Your password must contain set of numbers, letters (both Capital and small) and special symbols.',bg='#ffffff',font='-size 18')
    fpaes3.place(x=0,y=85)

    Ques1 = Label(tac,background='white',text='Q1) What is your Nickname? (All words must be in capital letters)',font='-size 18')
    Ques1.place(x=0,y=125)

    Answ1 = Entry(tac,background='#ffffff',textvariable=aone,font='-size 17',width=61)
    Answ1.place(x=710,y=127)
    Answ1.insert(0,AONE)

    Ques2 = Label(tac,background='white',text='Q2) Which city you love the most? (All words must be in capital letters)',font='-size 18')
    Ques2.place(x=0,y=165)

    Answ2 = Entry(tac,background='#ffffff',textvariable=atwo,font='-size 17',width=57)
    Answ2.place(x=764,y=167)
    Answ2.insert(0,ATWO)

    Ques3 = Label(tac,background='white',text='Q3) What is the name of your first school? (All words must be in capital letters)',font='-size 18')
    Ques3.place(x=0,y=205)

    Answ3 = Entry(tac,background='#ffffff',textvariable=athr,font='-size 17',width=50)
    Answ3.place(x=855,y=207)
    Answ3.insert(0,ATHR)

    CanceL = Button(changeframe,text='Cancel',cursor='hand2',background='#006F50',foreground='#ffffff',font='-size 17',height=1,width=10,command=gOback)
    CanceL.place(x=5,y=596)

    ConFirm = Button(changeframe,text='confirm',cursor='hand2',background='#006F50',foreground='#ffffff',font='-size 17',height=1,width=10,command=changerec)
    ConFirm.place(x=1390,y=596)

def loggedin(N,G,Do,Ph,Ei,R,Ex,C,Dg,U,Pa,Ao,At,Ar):
    def cnvpcall():
        cnvp(N,G,Do,Ph,Ei,R,Ex,C,Dg,U,Pa,Ao,At,Ar)

    def LogOut():
        signinfn()

    def ChAngE():
        ChangE(N,G,Do,Ph,Ei,R,Ex,C,Dg,U,Pa,Ao,At,Ar)
        
    loggedinframe = Frame(window,bg='#ffffff',height=649,width=1536,background='#ffffff')
    loggedinframe.place(x=0,y=191)

    userinfo = LabelFrame(loggedinframe,text='User Information:',font='-size 18 -weight bold',bd=6,bg='#ffffff',height=165,width=1535)
    userinfo.place(x=0,y=5)
    
    helLO = Label(userinfo,background='white',text='Hello:',font='-size 18')
    helLO.place(x=5,y=5)

    DrName = Label(userinfo,background='white',text=N,font='-size 18',fg='#006f50')
    DrName.place(x=75,y=5)

    regNO = Label(userinfo,background='white',text='Registeration Number:',font='-size 18')
    regNO.place(x=1150,y=5)

    REgNo = Label(userinfo,background='white',text=R,font='-size 18',fg='#006f50')
    REgNo.place(x=1400,y=5)

    DrDgr = Label(userinfo,background='white',text=Dg,font='-size 18',fg='#006f50')
    DrDgr.place(x=5,y=45)

    Phnolbl = Label(userinfo,background='white',text='Phone Number:',font='-size 18')
    Phnolbl.place(x=1150,y=45)

    PhnNo = Label(userinfo,background='white',text=Ph,font='-size 18',fg='#006f50')
    PhnNo.place(x=1330,y=45)

    DrDgr = Label(userinfo,background='white',text=C,font='-size 18',fg='#006f50')
    DrDgr.place(x=5,y=85)

    chngbtn = Button(userinfo,text='Change',cursor='hand2',background='#006F50',foreground='#ffffff',font='-size 15',height=0,width=8,command=ChAngE)
    chngbtn.place(x=1150,y=85)

    LogouT = Button(userinfo,text='Sign Out',cursor='hand2',background='#006F50',foreground='#ffffff',font='-size 15',height=0,width=8,command=LogOut)
    LogouT.place(x=1422,y=85)

    CrtNwPrscrpn= Button(loggedinframe,text='Create\nNew\nVoice\nPrescription',cursor='hand2',background='#006F50',foreground='#ffffff',font='-size 18',height=4,width=10,command=cnvpcall)
    CrtNwPrscrpn.place(x=50,y=335)

    hst = sqlite3.connect('drtalk')
    hcr = hst.cursor()
    hcr.execute("""select * from ps where DrregNO=?""",[R])
    hst.commit()
    res = hcr.fetchall()
    hst.close()

    HistoryFrame = Listbox(loggedinframe,bg='#ffffff',font='-size 18')
    HistoryFrame.place(x=250,y=170,height=475,width=1283)

    scroLLbar = Scrollbar(HistoryFrame,jump=0)
    scroLLbar.pack(side=RIGHT,fill=Y)

    for row in res:
        NameVAR=row[0]
        GendVAR=row[1]
        AgeVAR=row[2]
        SymptVAR=row[3]
        DiagVAR=row[4]
        PrscVAR=row[5]
        AdvVAR=row[6]
        PrnoVAR=row[8]
        HistoryFrame.insert(END, f"Name:{NameVAR}    Gender:{GendVAR}    Age:{row[2]}    Document Name:{PrnoVAR}")
        HistoryFrame.insert(END,'')
    HistoryFrame.config(yscrollcommand=scroLLbar.set)
    scroLLbar.config(command=HistoryFrame.yview)

def FPass():
    usname = StringVar()
    Passwd = StringVar()
    Avn = StringVar()
    Tu = StringVar()
    Tre = StringVar()
    PasswdPattern=r'(?=.*\d)(?=.*[a-z])(?=.*[A-Z])(?=.*\W)'
    def cpdata():
        docdb = sqlite3.connect('drtalk')
        drcr = docdb.cursor()
        drcr.execute("""select * from drs where Uname=?""",[usname.get()])
        docdb.commit()
        res=drcr.fetchall()
        docdb.close()
        for row in res:
            AOne=row[11]
            ATwo=row[12]
            AThr=row[13]
        if AOne==Avn.get() and ATwo==Tu.get() and AThr==Tre.get() and re.match(PasswdPattern,Passwd.get()):
            docdb = sqlite3.connect('drtalk')
            drcr = docdb.cursor()
            drcr.execute("""UPDATE drs SET Passwd=? where Uname=?""",[Passwd.get(),usname.get()])
            docdb.commit()
            docdb.close()
            window.destroy()
        else:
            messagebox.showinfo("Wrong Entries", "Please Enter Informations in correct format")

    fpassframe = Frame(window,bg='#ffffff',height=649,width=1536,background='#ffffff')
    fpassframe.place(x=0,y=191)
    
    unam = Label(fpassframe,background='white',text='Username',font='-size 18 -weight bold')
    unam.place(x=590,y=50)

    usrname = Entry(fpassframe,textvariable=usname,background='#ffffff',font='-size 17',borderwidth=2,relief='groove')
    usrname.place(x=750,y=52)

    QaA = LabelFrame(fpassframe,text='Question and Answers:',font='-size 18 -weight bold',bd=6,bg='#ffffff',height=285,width=1535)
    QaA.place(x=0,y=130)

    Ques1 = Label(QaA,background='white',text='Q1) What is your Nickname? (All words must be in capital letters)',font='-size 18')
    Ques1.place(x=0,y=5)

    Answ1 = Entry(QaA,background='#ffffff',textvariable=Avn,font='-size 17',width=61)
    Answ1.place(x=710,y=5)

    Ques2 = Label(QaA,background='white',text='Q2) Which city you love the most? (All words must be in capital letters)',font='-size 18')
    Ques2.place(x=0,y=110)

    Answ2 = Entry(QaA,background='#ffffff',textvariable=Tu,font='-size 17',width=57)
    Answ2.place(x=764,y=110)

    Ques3 = Label(QaA,background='white',text='Q3) What is the name of your first school? (All words must be in capital letters)',font='-size 18')
    Ques3.place(x=0,y=205)

    Answ3 = Entry(QaA,background='#ffffff',textvariable=Tre,font='-size 17',width=50)
    Answ3.place(x=855,y=207)

    npaSS = Label(fpassframe,background='white',text='New Password:',font='-size 18 -weight bold')
    npaSS.place(x=570,y=450)

    npass = Entry(fpassframe,textvariable=Passwd,background='#ffffff',font='-size 17',borderwidth=2,relief='groove')
    npass.place(x=750,y=452)

    confirM = Button(fpassframe,text='Confirm',cursor='hand2',background='#006F50',foreground='#ffffff',font='-size 17',height=1,width=10,command=cpdata)
    confirM.place(x=600,y=570)

    SiGnin = Button(fpassframe,text='Back',cursor='hand2',background='#006F50',foreground='#ffffff',font='-size 17',height=1,width=10,command=signinfn)
    SiGnin.place(x=800,y=570)

def signinfn():
    uname = StringVar()
    passwd = StringVar()

    def FPaSs():
        FPass()
        
    def signin():
        docdb = sqlite3.connect('drtalk')
        drcr = docdb.cursor()
        drcr.execute("""select * from drs where Uname=?""",[uname.get()])
        docdb.commit()
        res=drcr.fetchall()
        docdb.close()
        for row in res:
            NAme=row[0]
            GNdr=row[1]
            DOb=row[2]
            PHno=row[3]
            EId=row[4]
            RGno=row[5]
            EXp=row[6]
            CLn=row[7]
            DGr=row[8]
            UName=row[9]
            PAsswd=row[10]
            AOne=row[11]
            ATwo=row[12]
            AThr=row[13]
        try:
            if PAsswd == passwd.get():
                loggedin(NAme,GNdr,DOb,PHno,EId,RGno,EXp,CLn,DGr,UName,PAsswd,AOne,ATwo,AThr)
        except UnboundLocalError:
            messagebox.showinfo("Unknown Entry", "No records available at this username.\nPlease get registerd.")
        
    signinframe = Frame(window,bg='#ffffff',height=649,width=1536,background='#ffffff')
    signinframe.place(x=0,y=191)
    
    unam = Label(signinframe,background='white',text='Username',font='-size 18 -weight bold')
    unam.place(x=590,y=30)

    usrname = Entry(signinframe,background='#ffffff',textvariable=uname,font='-size 17',borderwidth=2,relief='groove')
    usrname.place(x=750,y=30)

    pas = Label(signinframe,background='white',text='Password',font='-size 18 -weight bold')
    pas.place(x=590,y=100)

    paswd = Entry(signinframe,background='#ffffff',textvariable=passwd,font='-size 17',borderwidth=2,relief='groove',show="*")
    paswd.place(x=750,y=100)

    sgnin = Button(signinframe,text='Sign in',cursor='hand2',background='#006F50',foreground='#ffffff',font='-size 17',height=1,width=10,command=signin)
    sgnin.place(x=640,y=170)

    sgnup = Button(signinframe,text='Sign up',cursor='hand2',background='#006F50',foreground='#ffffff',font='-size 17',height=1,width=10,command=signupfn)
    sgnup.place(x=805,y=170)

    fpass = Label(signinframe,background='white',text='Forgot password ~~>',font='-size 14')
    fpass.place(x=590,y=245)

    fpbtn = Button(signinframe,text='Click Here',cursor='hand2',background='#006F50',foreground='#ffffff',font='-size 14',height=1,width=15,command=FPaSs)
    fpbtn.place(x=805,y=240)

def signupfn():
    fname = StringVar()
    gendr = StringVar()
    dob = StringVar()
    phoneno = StringVar()
    eid = StringVar()
    regno = StringVar()
    exp = StringVar()
    cln = StringVar()
    dgr = StringVar()
    uname = StringVar()
    passwd = StringVar()
    aone = StringVar()
    atwo = StringVar()
    athr = StringVar()

    def signup():
        nmatch=0
        gmatch=0
        dobmatch=0
        phnomatch=0
        eidmatch=0
        regnomatch=0
        expmatch=0
        dgrmatch=0
        passwdmatch=0
        aonematch=0
        atwomatch=0
        athrmatch=0
        NamePattern=r'[a-zA-Z]|\.'
        if re.match(NamePattern,fname.get()):
            nmatch=1
        else:
            nmatch=0
        GenderPattern=r'Male|Female|Others'
        if re.match(GenderPattern,gendr.get()):
            gmatch=1
        else:
            gmatch=0
        DobPattern=r'\d{2}/\d{2}/\d{4}'
        if re.match(DobPattern,dob.get()):
            dobmatch=1
        else:
            dobmatch=0
        PhnoPattern=r'\d{10}'
        if re.match(PhnoPattern,phoneno.get()):
            phnomatch=1
        else:
            phnomatch=0
        EidPattern=r'[a-zA-Z0-9.]*@[a-zA-Z]*\.com'
        if re.match(EidPattern,eid.get()):
            eidmatch=1
        else:
            eidmatch=0
        RegnoPattern=r'\d{6}'
        if re.match(RegnoPattern,regno.get()):
            regnomatch=1
        else:
            regnomatch=0
        ExpPattern=r'\d{2}'
        if re.match(ExpPattern,exp.get()):
            expmatch=1
        else:
            expmatch=0
        DgrPattern=r'[a-zA-Z0-9. -]'
        if re.match(DgrPattern,dgr.get()):
            dgrmatch=1
        else:
            dgrmatch=0
        PasswdPattern=r'(?=.*\d)(?=.*[a-z])(?=.*[A-Z])(?=.*\W)'
        if re.match(PasswdPattern,passwd.get()):
            passwdmatch=1
        else:
            passwdmatch=0
        AnsPattern=r'[A-Z]+'
        if re.match(AnsPattern,aone.get()):
            aonematch=1
        else:
            aonematch=0
        if re.match(AnsPattern,atwo.get()):
            atwomatch=1
        else:
            atwomatch=0
        if re.match(AnsPattern,athr.get()):
            athrmatch=1
        else:
            athrmatch=0
        if nmatch and gmatch and dobmatch and phnomatch and eidmatch and regnomatch and expmatch and dgrmatch and passwdmatch and aonematch and atwomatch and athrmatch==1:
            print('Matched')
            docdb = sqlite3.connect('drtalk')
            drcr = docdb.cursor()
            drcr.execute("""create table if not exists drs(Name VARCHAR,Gender VARCHAR,Dob VARCHAR,PhoneNo VARCHAR UNIQUE,Eid VARCHAR UNIQUE,RegNo VARCHAR UNIQUE,Exp VARCHAR,Cln VARCHAR,Dgr VARCHAR,Uname VARCHAR PRIMARY KEY,Passwd VARCHAR,Aone VARCHAR,Atwo VARCHAR,Athr VARCHAR)""")
            drcr.execute("""insert into drs VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",(fname.get(), gendr.get(), dob.get(), phoneno.get(), eid.get(), regno.get(), exp.get(), cln.get(), dgr.get(), uname.get(), passwd.get(), aone.get(), atwo.get(), athr.get()))
            docdb.commit()
            docdb.close()
            signinfn()
        else:
            print('Not Matched')
            messagebox.showinfo("Wrong Entries", "Please Enter Informations in correct format")
    
    signupframe = Frame(window,bg='#ffffff',height=649,width=1536,background='#ffffff')
    signupframe.place(x=0,y=191)
    
    personalinfo = LabelFrame(signupframe,text='Personal Information:',font='-size 18 -weight bold',bd=6,bg='#ffffff',height=130,width=1535)
    personalinfo.place(x=0,y=5)

    nam = Label(personalinfo,text='Fullname:',bg='#ffffff',font='-size 18')
    nam.place(x=5,y=5)

    nme = Entry(personalinfo,bg='#ffffff',textvariable=fname,font='-size 17',width=107)
    nme.place(x=119,y=7)

    gndr = Label(personalinfo,text='Gender:',bg='#ffffff',font='-size 18')
    gendrbtn1 = Radiobutton(personalinfo,text='Male',tristatevalue='x',bg='#ffffff',variable=gendr,value='Male',font='-size 16')
    gendrbtn2 = Radiobutton(personalinfo,text='Female',tristatevalue='x',bg='#ffffff',variable=gendr,value='Female',font='-size 16')
    gendrbtn3 = Radiobutton(personalinfo,text='Others',tristatevalue='x',bg='#ffffff',variable=gendr,value='Others',font='-size 16')
    gndr.place(x=5,y=45)
    gendrbtn1.place(x=100,y=44)
    gendrbtn2.place(x=180,y=44)
    gendrbtn3.place(x=286,y=44)

    DoB = Label(personalinfo,text='D.O.B (dd/mm/yyyy):',bg='#ffffff',font='-size 18')
    DoB.place(x=395,y=45)
    Dob = Entry(personalinfo,bg='#ffffff',textvariable=dob,font='-size 17',width=9)
    Dob.place(x=628,y=47)

    Phn = Label(personalinfo,text='Contact Number:',bg='#ffffff',font='-size 18')
    Phn.place(x=760,y=45)
    Phno = Entry(personalinfo,bg='#ffffff',textvariable=phoneno,font='-size 17',width=10)
    Phno.place(x=950,y=47)

    Eml = Label(personalinfo,text='Email Id:',bg='#ffffff',font='-size 18')
    Eml.place(x=1110,y=45)
    Email = Entry(personalinfo,bg='#ffffff',textvariable=eid,font='-size 17',width=23)
    Email.place(x=1210,y=47)

    professionalinfo = LabelFrame(signupframe,text='Professional Information:',font='-size 18 -weight bold',bd=6,bg='#ffffff',height=130,width=1535)
    professionalinfo.place(x=0,y=140)

    Regno = Label(professionalinfo,text='Registeration Number:',bg='#ffffff',font='-size 18')
    Regno.place(x=0,y=5)

    RegNo = Entry(professionalinfo,bg='#ffffff',textvariable=regno,font='-size 17',width=6)
    RegNo.place(x=250,y=7)

    ExP = Label(professionalinfo,text='Experience (In Years):',bg='#ffffff',font='-size 18')
    ExP.place(x=350,y=5)

    Exp = Entry(professionalinfo,bg='#ffffff',textvariable=exp,font='-size 17',width=2)
    Exp.place(x=595,y=7)

    ClN = Label(professionalinfo,text='Clinic Name:',bg='#ffffff',font='-size 18')
    ClN.place(x=644,y=5)

    Cln = Entry(professionalinfo,bg='#ffffff',textvariable=cln,font='-size 17',width=55)
    Cln.place(x=793,y=7)

    DgR = Label(professionalinfo,text='Education Qualification:',bg='#ffffff',font='-size 18')
    DgR.place(x=0,y=45)

    Dgr = Entry(professionalinfo,bg='#ffffff',textvariable=dgr,font='-size 17',width=96)
    Dgr.place(x=260,y=47)

    UsrNme = Label(signupframe,text='Username:',font='-size 18 -weight bold',bg='#ffffff')
    UsrNme.place(x=4,y=270)

    usernam = Entry(signupframe,background='#ffffff',textvariable=uname,font='-size 17',width=48)
    usernam.place(x=130,y=274)

    pas = Label(signupframe,background='white',text='Password:',font='-size 18 -weight bold')
    pas.place(x=760,y=270)

    psswd = Entry(signupframe,background='#ffffff',textvariable=passwd,font='-size 17',width=48)
    psswd.place(x=890,y=274)

    tac = LabelFrame(signupframe,text='Terms and Condition:',font='-size 18 -weight bold',bd=6,bg='#ffffff',height=285,width=1535)
    tac.place(x=0,y=307)

    fpaes1 = Label(tac,text='By chance if you forgot your password then by clicking on forgot password, You have answer the below given questions and you can change',bg='#ffffff',font='-size 18')
    fpaes1.place(x=0,y=5)

    fpaes2 = Label(tac,text='your password. Another thing is the application will accept only those Email Ids which are based on .com servers so you must enter only tho-',bg='#ffffff',font='-size 18')
    fpaes2.place(x=0,y=45)

    fpaes3 = Label(tac,text='-se emails having .com at end. Your password must contain set of numbers, letters (both Capital and small) and special symbols.',bg='#ffffff',font='-size 18')
    fpaes3.place(x=0,y=85)

    Ques1 = Label(tac,background='white',text='Q1) What is your Nickname? (All words must be in capital letters)',font='-size 18')
    Ques1.place(x=0,y=125)

    Answ1 = Entry(tac,background='#ffffff',textvariable=aone,font='-size 17',width=61)
    Answ1.place(x=710,y=127)

    Ques2 = Label(tac,background='white',text='Q2) Which city you love the most? (All words must be in capital letters)',font='-size 18')
    Ques2.place(x=0,y=165)

    Answ2 = Entry(tac,background='#ffffff',textvariable=atwo,font='-size 17',width=57)
    Answ2.place(x=764,y=167)

    Ques3 = Label(tac,background='white',text='Q3) What is the name of your first school? (All words must be in capital letters)',font='-size 18')
    Ques3.place(x=0,y=205)

    Answ3 = Entry(tac,background='#ffffff',textvariable=athr,font='-size 17',width=50)
    Answ3.place(x=855,y=207)

    preV = Button(signupframe,text='Sign In',cursor='hand2',background='#006F50',foreground='#ffffff',font='-size 17',height=1,width=10,command=signinfn)
    preV.place(x=5,y=596)

    Prev = Label(signupframe,background='white',text='(Previous)',font='-size 18')
    Prev.place(x=150,y=601)

    Nxt = Button(signupframe,text='Sign Up',cursor='hand2',background='#006F50',foreground='#ffffff',font='-size 17',height=1,width=10,command=signup)
    Nxt.place(x=1390,y=596)

    nxT = Label(signupframe,background='white',text='(Next)',font='-size 18')
    nxT.place(x=1315,y=601)    

signinfn()

window.mainloop()
